# -*- coding: utf-8 -*-
"""Build VB2026_WebDev_IDUBmandus.docx.

Format per guidebook: A4, margin 4-4-3-3 cm, Times New Roman 12pt, spasi 1,5,
maksimal 13 halaman isi (sampul dan lampiran tidak dihitung).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\FinancialTracker\docs\veternity\VB2026_WebDev_IDUBmandus.docx"
IMG = r"D:\FinancialTracker\docs\veternity\img"

ACCENT = RGBColor(0x0B, 0x4F, 0x7A)
doc = Document()

sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.top_margin, sec.left_margin = Cm(4), Cm(4)
sec.bottom_margin, sec.right_margin = Cm(3), Cm(3)

st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(12)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = st.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(4)
pf.space_before = Pt(0)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _fmt(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return run


def para(text="", size=12, bold=False, italic=False, align=None, space_after=4,
         space_before=0, spacing=1.5, color=None):
    p = doc.add_paragraph()
    pfm = p.paragraph_format
    if spacing == 1.5:
        pfm.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    else:
        pfm.line_spacing = spacing
    pfm.space_after = Pt(space_after)
    pfm.space_before = Pt(space_before)
    pfm.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        _fmt(p.add_run(text), size, bold, italic, color)
    return p


def rich(parts, size=12, align=None, space_after=4):
    p = doc.add_paragraph()
    pfm = p.paragraph_format
    pfm.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pfm.space_after = Pt(space_after)
    pfm.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    for t, b, i in parts:
        _fmt(p.add_run(t), size, b, i)
    return p


_num = [0]


def h1(text):
    _num[0] = 0
    return para(text, size=13, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=10, space_after=5, color=ACCENT)


def h2(text):
    _num[0] = 0
    return para(text, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=7, space_after=4)


def h3(text):
    return para(text, size=12, bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=6, space_after=3)


def _listp(style, text, size=12):
    p = doc.add_paragraph(style=style)
    pfm = p.paragraph_format
    pfm.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pfm.space_after = Pt(2)
    pfm.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _fmt(p.add_run(text), size)
    return p


def bullet(text, size=12):
    return _listp("List Bullet", text, size)


def numbered(text, size=12):
    """Manual numbering: python-docx shares one list instance across the doc,
    so the built-in List Number style would keep counting past each section."""
    _num[0] += 1
    p = doc.add_paragraph()
    pfm = p.paragraph_format
    pfm.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pfm.space_after = Pt(2)
    pfm.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pfm.left_indent = Cm(0.75)
    pfm.first_line_indent = Cm(-0.75)
    _fmt(p.add_run(f"{_num[0]}.  {text}"), size)
    return p


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def table(rows, widths=None, size=9.5, header=True):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cell = cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            pfm = p.paragraph_format
            pfm.line_spacing = 0.95
            pfm.space_after = Pt(0.5)
            pfm.space_before = Pt(0.5)
            pfm.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _fmt(p.add_run(val), size, bold=(header and ri == 0))
            if widths:
                cell.width = Cm(widths[ci])
        if header and ri == 0:
            for c in cells:
                _shade(c, "DCE6F1")
            _repeat_header(t.rows[0])
    return t


def figure(fname, caption, width=9.0):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    p.add_run().add_picture(os.path.join(IMG, fname), width=Cm(width))
    para(caption, size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=7, spacing=1.0)


def page_break(p):
    """Break before p, so no stray empty paragraph can land on its own page."""
    p.paragraph_format.page_break_before = True
    return p


# ====================================================================== COVER
para("", space_after=0)
para("PROPOSAL KARYA", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Diajukan untuk Babak Penyisihan", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Web Development Competition, Veternity Beraksi 2026", bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
para("Tema: “Bridging the Gap: Digital Platforms for Equitable Economic Access”",
     size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Sub Tema: Micro-Capital Crowdfunding & Financial Management",
     size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para("TRACR", size=30, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, color=ACCENT)
para("Buku Kas Digital yang Tetap Berjalan Saat Sinyal Mati: Platform Manajemen Keuangan Mikro "
     "Berbasis Offline-First untuk Pelaku Usaha Mikro dan Masyarakat Unbanked",
     size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=26)
table([
    ["Nama Tim", "IDUB mandus"],
    ["Anggota 1", "Muhammad Choirudin Ammar (25/556251/TK/62735)"],
    ["Anggota 2", "Muhammad Raihan Surya (25/560713/TK/63338)"],
    ["Anggota 3", "Ahmad Rafi Firdaus (25/560526/TK/63314)"],
    ["Perguruan Tinggi", "Universitas Gadjah Mada"],
    ["Program Studi", "Teknologi Informasi, Fakultas Teknik"],
    ["Tautan Website", "https://tracr-ai.vercel.app"],
    ["Repositori GitHub", "https://github.com/astrorehan/FinancialTracker"],
], widths=[4.2, 9.8], size=11, header=False)
para("2026", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=16)

# ======================================================= 1. JUDUL & NAMA TIM
page_break(h1("1. JUDUL KARYA & NAMA TIM"))
rich([("Judul Karya: ", True, False),
      ("Tracr, Buku Kas Digital yang Tetap Berjalan Saat Sinyal Mati: Platform Manajemen "
       "Keuangan Mikro Berbasis ", False, False), ("Offline-First", False, True),
      (" untuk Pelaku Usaha Mikro dan Masyarakat ", False, False), ("Unbanked", False, True),
      (".", False, False)])
rich([("Nama Tim: ", True, False), ("IDUB mandus, Universitas Gadjah Mada.", False, False)])
rich([("Sub Tema: ", True, False),
      ("Sub Tema 2, Micro-Capital Crowdfunding & Financial Management.", False, False)])
para("Tracr adalah aplikasi web pencatatan keuangan mikro yang berfungsi penuh tanpa koneksi "
     "internet, menggunakan bahasa sehari-hari alih-alih istilah akuntansi, dan dapat dipasang ke "
     "layar utama telepon genggam tanpa melalui toko aplikasi. Satu antarmuka melayani dua "
     "kebutuhan sekaligus: pencatatan keuangan pribadi dan pembukuan usaha mikro (Buku Usaha), "
     "yang dipisahkan secara tegas melalui mekanisme multi-book.")

# ================================================== 2. LATAR BELAKANG MASALAH
h1("2. LATAR BELAKANG MASALAH")
h2("2.1 Kesenjangan Ekonomi yang Disasar")
para("UMKM adalah tulang punggung perekonomian Indonesia: sekitar 64,2 juta unit yang menyumbang "
     "61,07% Produk Domestik Bruto dan menyerap 97% tenaga kerja nasional (Kementerian Koperasi "
     "dan UKM, 2021). Strukturnya timpang jauh ke bawah: basis data tunggal SIDT-UMKM mencatat "
     "30.209.069 unit usaha terverifikasi per 31 Desember 2025 dan 99,70% di antaranya berskala "
     "mikro (Kementerian UMKM, 2025), yakni warung kelontong, pedagang kaki lima, penjual makanan "
     "rumahan, dan penjahit. Justru pada kelompok itulah pencatatan keuangan paling jarang "
     "ditemukan. Kajian lapangan di Pekanbaru mendapati pelaku UMKM umumnya hanya mencatat "
     "pemasukan dan pengeluaran seadanya, bahkan ada yang tidak mencatat sewa tempat sebagai beban "
     "sehingga usahanya tampak berlaba besar (Fatwitawati, 2018); dari 30 pelaku UMKM di Medan "
     "Timur, hanya sebagian kecil yang memahami akuntansi sederhana (Fadhia dan Ningsih, 2024).")
para("Akses keuangan formal pun belum merata. SNLIK 2025 mencatat indeks inklusi keuangan 80,51% "
     "berbanding indeks literasi keuangan 66,46%, sehingga sekitar satu dari lima penduduk dewasa "
     "masih di luar layanan keuangan formal dan selisih 14,05 poin itu memperlihatkan banyak orang "
     "memegang produk keuangan tanpa memahaminya; di perdesaan jaraknya melebar menjadi 75,70% "
     "berbanding 59,60% (OJK dan BPS, 2025). Transaksi kelompok ini karenanya berjalan tunai, dan "
     "solusi digital yang bertumpu pada sinkronisasi rekening bank (open banking, agregasi mutasi, "
     "kategorisasi otomatis dari pesan bank) tidak dapat menjangkau mereka karena data sumbernya "
     "memang tidak pernah ada.")
para("Kesenjangan itu diperparah sebaran kualitas jaringan yang timpang. Pendataan Potensi Desa "
     "2024, sensus atas 84.276 wilayah setingkat desa, mencatat 17,52% desa hanya menerima sinyal "
     "seluler lemah dan 3,70% tidak menerima sinyal sama sekali; 2,47% desa masih bergantung pada "
     "internet 2G/EDGE dan 1,26% tanpa sinyal internet (BPS, 2024). Bagi pedagang pasar yang "
     "sedang melayani pembeli, aplikasi yang menampilkan layar kosong belasan detik karena "
     "menunggu peladen bukan sekadar tidak nyaman; aplikasi itu tidak dapat dipakai.")

h2("2.2 Empat Lapis Hambatan")
table([
    ["Hambatan", "Wujud nyata di lapangan", "Konsekuensi ekonomi"],
    ["Geografis", "Sinyal putus-putus atau hilang di pasar, pelosok desa, dan area produksi",
     "Pencatatan tertunda lalu terlupakan; data keuangan usaha tidak pernah lengkap"],
    ["Literasi", "Istilah “aset”, “liabilitas”, “rekonsiliasi”, “jurnal” pada aplikasi yang ada",
     "Pengguna awam merasa aplikasi bukan untuk dirinya dan berhenti di langkah pertama"],
    ["Perangkat & kuota", "Telepon kelas pemula, memori hampir penuh, kuota terbatas",
     "Enggan memasang aplikasi puluhan megabita dari toko aplikasi"],
    ["Akses permodalan", "Tidak ada rekam jejak keuangan tertulis saat mengajukan KUR atau pinjaman koperasi",
     "Pengajuan modal ditolak; usaha tidak naik kelas meskipun sehat secara arus kas (IFC, 2025)"],
], widths=[2.6, 5.7, 5.7])
para("Hambatan keempat adalah simpul persoalan yang sesungguhnya, dan menjadi alasan karya ini "
     "relevan dengan equitable economic access, bukan sekadar financial management. Pelaku usaha "
     "mikro dapat memiliki arus kas sehat bertahun-tahun namun tetap ditolak saat mengajukan "
     "modal, semata-mata karena tidak memiliki dokumen yang membuktikannya. IFC menyebut akarnya "
     "sebagai asimetri informasi: pemberi pinjaman tidak memiliki cukup keterangan untuk menilai "
     "kelayakan kredit UMKM, sehingga risikonya dipersepsikan lebih tinggi (IFC, 2025). Bagi "
     "lembaga keuangan, usaha tanpa catatan adalah usaha yang tidak terlihat.", space_before=4)

h2("2.3 Mengapa Solusi yang Ada Belum Menjawab")
para("Aplikasi pembukuan yang beredar mengandaikan tiga hal yang justru tidak dimiliki kelompok "
     "sasaran: koneksi stabil, keakraban dengan istilah akuntansi, dan kesediaan memasang aplikasi "
     "berukuran besar. Sebagian juga mengunci laporan laba rugi dan ekspor data di balik langganan "
     "berbayar, tepat ketika data itu mulai berguna. Yang paling menentukan, banyak aplikasi "
     "memperlakukan dukungan luring hanya sebagai penyimpanan sementara untuk membaca data; ketika "
     "pengguna mencoba menulis, operasi tersebut gagal atau menggantung. Padahal justru menulis "
     "itulah inti sebuah buku kas. Evaluasi aplikasi akuntansi UMKM berbasis awan menyimpulkan "
     "kunci penerimaan pengguna justru terletak pada logika debit-kredit yang disembunyikan di "
     "balik antarmuka (Mayasari dkk., 2025).")

h2("2.4 Rumusan Masalah")
numbered("Bagaimana merancang platform manajemen keuangan mikro yang tetap menerima pencatatan "
         "baru tanpa koneksi internet, dan menyinkronkannya secara utuh serta berurutan ketika "
         "koneksi pulih?")
numbered("Bagaimana menyajikan konsep modal, laba, piutang, dan utang kepada pengguna tanpa latar "
         "belakang akuntansi, tanpa mengorbankan ketepatan perhitungan?")
numbered("Bagaimana mengubah kebiasaan mencatat harian menjadi rekam jejak keuangan yang dapat "
         "dipertanggungjawabkan sebagai dokumen pendukung akses permodalan formal?")

# ============================================ 3. SOLUSI, TUJUAN, DAN MANFAAT
h1("3. SOLUSI, TUJUAN, DAN MANFAAT APLIKASI")
h2("3.1 Gagasan Solusi")
para("Tracr menjawab ketiga rumusan masalah melalui tiga pilar rancangan yang seluruhnya sudah "
     "terimplementasi dan berjalan pada website yang dinilai.")
h3("Pilar 1: Offline-First yang Sesungguhnya")
para("Tracr tidak sekadar menyimpan salinan data untuk dibaca. Setiap operasi penulisan yang "
     "dilakukan saat perangkat luring masuk ke antrian mutasi di IndexedDB peramban, dengan "
     "cadangan localStorage. Ketika koneksi pulih, pekerja antrian memutar ulang mutasi satu per "
     "satu sesuai urutan masuk (FIFO) ke basis data. Rancangan ini menangani tiga persoalan yang "
     "biasanya diabaikan implementasi luring sederhana:")
numbered("Ketergantungan antarmutasi. Bila pengguna membuat produk baru saat luring lalu langsung "
         "menjualnya, mutasi penjualan merujuk identitas sementara yang belum ada di peladen. "
         "Fungsi remapQueuedTempIds menukar seluruh kemunculan identitas itu dengan UUID asli "
         "begitu mutasi induknya tersimpan, sehingga rantai mutasi tidak putus.")
numbered("Kegagalan yang tidak menghilang diam-diam. Mutasi gagal diulang paling banyak tiga "
         "kali, lalu dipindahkan ke antrian gagal yang ditampilkan kepada pengguna beserta pesan "
         "galat dan tombol coba lagi. Data pengguna tidak pernah hilang tanpa pemberitahuan.")
numbered("Kejelasan status. Spanduk status menampilkan kondisi luring dan jumlah catatan yang "
         "menunggu sinkronisasi, sehingga pengguna tahu catatannya aman meskipun belum terkirim.")
para("Lebih dari 40 jenis operasi didukung secara luring: transaksi, akun, kategori, anggaran, "
     "target tabungan, utang-piutang, produk, dan cicilan.")
h3("Pilar 2: Bahasa Manusia, Bukan Bahasa Akuntansi")
para("Antarmuka berbahasa Indonesia secara bawaan dan sengaja menghindari istilah akuntansi. "
     "Pengguna tidak pernah bertemu kata “liabilitas”; yang muncul adalah jenis akun “Kartu "
     "Kredit” atau “Pinjaman” yang menjelaskan dirinya sebagai uang yang harus dibayar. Pada modul "
     "utang-piutang pilihannya “Pelanggan ngutang” dan “Saya ngutang”; pada modul usaha kolomnya "
     "bernama “Modal” dan hasilnya “Untung”. Rinciannya dibahas pada Bagian 8.")
h3("Pilar 3: Dua Buku dalam Satu Aplikasi")
para("Nasihat pertama setiap pendamping UMKM adalah memisahkan uang pribadi dari uang usaha. "
     "Tracr menjadikannya satu ketukan: setiap catatan dipartisi berdasarkan book_id, dan setiap "
     "buku berjenis personal atau business. Saat pengguna berpindah ke buku usaha, aplikasi "
     "membuka tiga alat yang tidak relevan bagi pengguna pribadi: Produk & Kasir Sederhana, "
     "Utang-Piutang, dan Laba Rugi.")
h3("Jembatan Menuju Akses Permodalan")
para("Ketiga pilar bermuara pada satu keluaran yang menjawab hambatan keempat pada Bagian 2.2. "
     "Setelah beberapa bulan mencatat, pengguna dapat menghasilkan Laporan Laba Rugi dan riwayat "
     "transaksi lengkap dalam bentuk PDF maupun CSV, dokumen yang selama ini tidak dimiliki "
     "pelaku usaha mikro saat berhadapan dengan bank, koperasi, atau program pendanaan. Tracr "
     "karena itu bukan sekadar alat catat-mencatat, melainkan pembentuk rekam jejak yang mengubah "
     "usaha tak terlihat menjadi usaha yang dapat dinilai.")

h2("3.2 Tujuan")
numbered("Membangun platform manajemen keuangan mikro yang berfungsi penuh, termasuk operasi "
         "penulisan, tanpa koneksi internet, dan menyinkronkan data secara berurutan serta dapat "
         "diaudit ketika koneksi pulih.")
numbered("Menyediakan antarmuka berbahasa Indonesia bebas istilah akuntansi sehingga dapat "
         "dipakai pelaku usaha mikro tanpa pelatihan formal.")
numbered("Memisahkan pembukuan usaha dari keuangan pribadi melalui mekanisme multi-book tanpa "
         "mengharuskan pengguna membuat akun kedua.")
numbered("Menghasilkan laporan laba rugi dan riwayat transaksi yang dapat diekspor sebagai "
         "dokumen pendukung pengajuan permodalan, serta menjaga seluruh fitur inti tetap gratis "
         "di atas arsitektur yang biayanya tidak bertambah linier terhadap jumlah pengguna.")

h2("3.3 Manfaat")
rich([("Bagi pelaku usaha mikro: ", True, False),
      ("mengetahui laba yang sebenarnya, bukan sekadar jumlah uang di laci; mengetahui siapa "
       "berutang dan sudah berapa lama; memiliki dokumen keuangan saat mengajukan modal; "
       "memisahkan uang usaha dari uang belanja rumah tangga.", False, False)])
rich([("Bagi pekerja informal dan masyarakat unbanked: ", True, False),
      ("memperoleh alat pencatatan yang tidak mensyaratkan rekening bank, koneksi tetap, maupun "
       "telepon genggam kelas atas.", False, False)])
rich([("Bagi ekosistem keuangan: ", True, False),
      ("bertambahnya pelaku usaha mikro yang memiliki rekam jejak tertulis memperbesar populasi "
       "calon penerima kredit yang dapat dinilai, sehingga menurunkan hambatan penyaluran "
       "pembiayaan mikro.", False, False)])

# ================================================ 4. RUANG LINGKUP & BATASAN
h1("4. RUANG LINGKUP & BATASAN APLIKASI")
h2("4.1 Ruang Lingkup")
table([
    ["Aspek", "Keterangan"],
    ["Platform & pengguna", "Aplikasi web berbasis peramban yang dapat dipasang sebagai Progressive Web App ke layar utama telepon genggam, untuk pelaku usaha mikro, pekerja informal, dan individu yang mengelola keuangan secara tunai"],
    ["Cakupan fungsional", "Pribadi: multi-akun, multi-mata uang, transaksi pemasukan/pengeluaran/transfer, kategori bertingkat, label, anggaran, tagihan berulang, target tabungan, cicilan, laporan. Usaha: produk & kasir sederhana, utang-piutang per pelanggan, laporan laba rugi, kontak pelanggan/pemasok"],
    ["Bahasa & mata uang", "Bahasa Indonesia (bawaan) dan Inggris; multi-mata uang dengan kurs yang dikelola pengguna dan dibekukan per transaksi"],
    ["Mode operasi", "Daring dan luring penuh (baca dan tulis); kanal tambahan berupa bot Telegram"],
], widths=[3.4, 10.6])
h2("4.2 Batasan")
table([
    ["Batasan", "Alasan rancangan"],
    ["Bukan dompet elektronik; tidak menyimpan uang dan tidak memproses pembayaran",
     "Aplikasi murni pencatatan sehingga tidak menimbulkan risiko dana pengguna; pemrosesan pembayaran juga memerlukan perizinan yang tidak dimiliki tim"],
    ["Tidak terhubung ke rekening bank (tanpa open banking); kurs dimasukkan manual",
     "Kelompok sasaran justru bertransaksi tunai, sehingga ketergantungan pada rekening akan menggugurkan relevansi solusi; kurs manual menjaga aplikasi tetap gratis (kolom source sudah disiapkan untuk pengisian otomatis kelak)"],
    ["Satu pengguna per buku (belum multi-kasir)",
     "Model keamanan bertumpu pada RLS auth.uid() = user_id; berbagi buku memerlukan model izin baru yang belum diuji cukup"],
    ["Manajemen stok belum tersedia",
     "Ditunda karena berisiko menimbulkan galat perhitungan pada demo langsung; produk menyimpan harga dan modal, belum jumlah persediaan"],
    ["Pengingat utang dikirim manual lewat tautan WhatsApp",
     "Pengiriman otomatis memerlukan verifikasi akun bisnis Meta yang masih dalam proses"],
], widths=[5.0, 9.0])
h2("4.3 Asumsi")
numbered("Pengguna memiliki telepon genggam dengan peramban modern yang mendukung Service Worker "
         "dan IndexedDB.")
numbered("Pengguna memiliki akses internet sesekali, tidak harus terus-menerus, untuk masuk "
         "pertama kali dan menyinkronkan data.")
numbered("Pengguna memiliki akun Google untuk masuk; opsi masuk tanpa Google merupakan "
         "pengembangan lanjutan yang sudah direncanakan.")

# ================================================= 5. TEKNOLOGI YANG DIPAKAI
h1("5. TEKNOLOGI YANG DIGUNAKAN")
h2("5.1 Ringkasan Tumpukan Teknologi")
table([
    ["Lapisan", "Teknologi", "Alasan pemilihan"],
    ["Frontend", "React 19 + TypeScript 6 + Vite 8 + Tailwind CSS v4",
     "Pemeriksaan tipe menyeluruh, waktu bangun cepat, keluaran statis murni, satu sumber token warna"],
    ["Data, rute, grafik", "TanStack Query v5, React Router 7, Recharts 3, React Hook Form 7 + Zod 4",
     "Singgahan permintaan dan pembaruan optimistis (titik sisip alami bagi antrian luring); rute dimuat malas; validasi berskema sebelum data dikirim"],
    ["PWA & luring", "vite-plugin-pwa + Workbox; IndexedDB + localStorage",
     "Service Worker, manifes, pemasangan ke layar utama; antrian mutasi, antrian gagal, dan singgahan kueri yang bertahan setelah aplikasi ditutup"],
    ["Basis data & autentikasi", "Supabase PostgreSQL (40 migrasi), Auth Google OAuth, Storage bucket privat",
     "Relasional dengan view, trigger, dan Row Level Security per pengguna; tanpa pengelolaan kata sandi sendiri; lampiran dilindungi URL bertanda tangan"],
    ["Sisi peladen", "Supabase Edge Functions (Deno) + pg_cron",
     "Tujuh fungsi terisolasi; kunci rahasia tidak pernah menyentuh peramban; penjadwal transaksi berulang"],
    ["Pengujian & penggelaran", "Vitest (187 kasus uji, 17 berkas); Vercel + Supabase",
     "Uji memusat pada logika perhitungan uang; keduanya berjalan pada paket gratis, aset diberi cache permanen sedangkan Service Worker tidak di-cache"],
], widths=[3.0, 4.6, 6.4])

h2("5.2 Fungsi Sisi Peladen dan Ketepatan Perhitungan Uang")
para("Tujuh Edge Function menangani seluruh pekerjaan yang tidak boleh berjalan di peramban. "
     "recurring-autopost dipanggil pg_cron setiap hari untuk memasang transaksi berulang yang "
     "jatuh tempo, menyusul periode terlewat, lalu memajukan tanggal jatuh tempo. tg-webhook dan "
     "wa-webhook mengubah pesan bot menjadi transaksi pada buku yang tertaut dan berbagi inti "
     "pemrosesan. send-push mengirim notifikasi pengingat tagihan, ai-analysis menyediakan "
     "ringkasan pengeluaran opsional (Bagian 6.4), sedangkan billing-checkout dan midtrans-webhook "
     "merupakan jalur pembayaran opsional yang tidak diaktifkan pada versi lomba.")
para("Seluruh nilai uang disimpan sebagai bilangan bulat satuan terkecil (bigint), yakni rupiah tanpa "
     "desimal, sen untuk mata uang dua desimal, satoshi untuk aset kripto. Konversi hanya terjadi "
     "di tepi sistem melalui src/lib/money.ts, sehingga galat pembulatan pecahan biner yang lazim "
     "pada aplikasi keuangan berbasis JavaScript tidak dapat muncul. Saldo akun pun tidak dihitung "
     "di sisi klien: view SQL account_balances menjumlahkan saldo awal ditambah seluruh mutasi "
     "bertanda, termasuk transfer yang mendebit akun asal dan mengkredit akun tujuan.",
     space_before=4)

# ============================================ 6. FITUR UTAMA & NILAI KEUNIKAN
h1("6. FITUR UTAMA & NILAI KEUNIKAN APLIKASI")
h2("6.1 Fitur Utama")
table([
    ["Kelompok", "Fitur"],
    ["Pencatatan inti",
     "Multi-akun (tunai, bank, kartu, dompet elektronik, kripto, saham, pinjaman) dengan saldo awal dan arsip; transaksi pemasukan/pengeluaran/transfer beserta kategori bertingkat berikon, label, penerima, dan lampiran struk; kalkulator pada kolom jumlah; transaksi terbagi ke beberapa kategori; penyaringan gabungan (akun, kategori, label, jenis, rentang tanggal, rentang jumlah, teks) yang dapat disimpan sebagai tampilan"],
    ["Perencanaan",
     "Anggaran per kategori atau menyeluruh dengan ambang peringatan; tagihan berulang dengan pemasangan otomatis; target tabungan; cicilan bertenor"],
    ["Laporan & data",
     "Pemasukan vs pengeluaran, komposisi kategori, tren kekayaan bersih, perbandingan antarperiode, peta panas harian; ekspor CSV dan PDF; impor CSV bervalidasi; cadangan dan pemulihan JSON"],
    ["Buku Usaha",
     "Produk & kasir sederhana (harga jual dan modal, keranjang sekali ketuk, satu penjualan = satu transaksi berikut rinciannya); utang-piutang dikelompokkan per orang dengan pembayaran sebagian, usia utang berbahasa manusia, dan pengingat WhatsApp; laporan Laba Rugi (Penjualan, Modal, Laba Kotor, Biaya, Laba Bersih) beserta produk terlaris"],
    ["Aksesibilitas",
     "Mode luring penuh untuk lebih dari 40 jenis operasi tulis; pemasangan PWA tanpa toko aplikasi; dwibahasa Indonesia-Inggris; pengaturan ukuran teks dan mode gelap; bot Telegram; notifikasi Web Push"],
], widths=[3.2, 10.8])

h2("6.2 Nilai Keunikan (Orisinalitas)")
para("Lima keputusan rekayasa berikut membedakan Tracr dari aplikasi pencatatan keuangan pada "
     "umumnya, dan seluruhnya dapat ditunjukkan pada kode sumber.")
rich([("1. Antrian mutasi luring dengan pemetaan ulang identitas sementara. ", True, False),
      ("Yang paling jarang ditemui bukan antriannya, melainkan penukaran identitas sementara "
       "dengan UUID asli pada seluruh mutasi yang masih mengantre, sehingga entitas yang dibuat "
       "dan langsung dipakai saat luring tidak menghasilkan rujukan menggantung.", False, False)])
rich([("2. Pembekuan harga jual dan modal pada setiap baris penjualan. ", True, False),
      ("Tabel transaction_items menyimpan unit_price dan unit_cost sebagai salinan saat penjualan "
       "terjadi, bukan rujukan ke harga produk saat ini. Ketika pedagang menaikkan harga bulan "
       "depan, laba bulan lalu tidak ikut berubah. Implementasi naif yang menggabungkan penjualan "
       "ke harga produk terkini menghasilkan laporan historis yang keliru, dan kekeliruan itu "
       "tidak terlihat sampai harga berubah.", False, False)])
rich([("3. Uang sebagai bilangan bulat satuan terkecil di seluruh sistem. ", True, False),
      ("Dari kolom basis data hingga hasil perhitungan laba tidak ada bilangan pecahan; ketepatan "
       "terjaga secara struktural, bukan karena pembulatan di akhir.", False, False)])
rich([("4. Saldo dihitung basis data, bukan klien. ", True, False),
      ("View account_balances adalah satu-satunya sumber kebenaran saldo, menghilangkan seluruh "
       "kelas galat “saldo di layar berbeda dengan saldo sebenarnya”.", False, False)])
rich([("5. Pembekuan kurs pada setiap transaksi. ", True, False),
      ("Setiap transaksi menyimpan base_amount dan fx_rate saat pembuatan, sehingga riwayat "
       "keuangan tidak berubah ketika kurs bergerak, perilaku yang benar secara akuntansi dan "
       "jarang diterapkan aplikasi sejenis.", False, False)])

h2("6.3 Keamanan dan Manajemen Data")
bullet("Row Level Security pada seluruh tabel dengan kebijakan auth.uid() = user_id; tabel "
       "penghubung menyalin user_id agar kebijakan tetap sederhana. Tidak ada jalur baca atau "
       "tulis yang dapat menembus batas antarpengguna, bahkan bila kunci publik klien bocor, "
       "sebab kunci tersebut memang dirancang untuk publik.")
bullet("Rahasia tidak pernah berada di peramban: seluruh kunci pihak ketiga berada pada Edge "
       "Functions, dan rahasia pemanggilan pg_cron disimpan pada tabel app_secrets yang terkunci "
       "RLS. Lampiran pun privat: bucket Storage tidak dapat diakses publik dan hanya terbuka "
       "melalui URL bertanda tangan berumur pendek.")
bullet("Validasi berlapis: Zod di sisi klien, batasan check di basis data (jumlah harus positif, "
       "arah utang harus salah satu dari dua nilai sah), dan RLS sebagai lapisan terakhir.")
bullet("Kedaulatan data: cadangan JSON menyeluruh dan ekspor CSV tersedia tanpa syarat.")

h2("6.4 Keberlanjutan Layanan")
para("Seluruh fitur inti (pencatatan, laporan, Buku Usaha, mode luring, dan ekspor data) "
     "direncanakan tetap gratis tanpa batas waktu; frontend berupa berkas statis sedangkan basis "
     "data dan autentikasi berjalan pada paket gratis. Untuk pembiayaan jangka panjang, Tracr "
     "menyiapkan lapisan analisis berbasis kecerdasan buatan sebagai layanan tambahan opsional "
     "berbasis kredit di luar jalur pencatatan, sehingga pengguna yang tidak pernah membelinya "
     "tetap memperoleh aplikasi yang utuh dan pengguna yang mampu menyubsidi yang tidak mampu. "
     "Pada versi yang dinilai dalam kompetisi ini, jalur pembayaran dinonaktifkan.")

# ================================================== 7. ARSITEKTUR & USE CASE
h1("7. ARSITEKTUR SISTEM & USE CASE DIAGRAM")
h2("7.1 Arsitektur Sistem")
para("Tracr menggunakan arsitektur serverless tiga lapis tanpa peladen aplikasi yang dikelola "
     "sendiri: frontend berupa berkas statis, seluruh logika istimewa pada Edge Functions, dan "
     "basis data yang menegakkan keamanannya sendiri melalui Row Level Security. Alur pencatatan "
     "luring dan sinkronisasi ulang digambarkan tersendiri pada Gambar 4 di Lampiran.")
figure("01-arsitektur.png", "Gambar 1. Arsitektur sistem Tracr")
h2("7.2 Rancangan Basis Data")
para("Basis data terdiri atas 40 migrasi bernomor. Seluruh tabel dipartisi per pengguna melalui "
     "user_id, dan tabel pembukuan juga dipartisi per buku melalui book_id.")
figure("03-erd.png", "Gambar 2. Rancangan basis data (inti yang relevan dengan sub tema)")
h2("7.3 Use Case Diagram")
figure("04-usecase.png", "Gambar 3. Use case diagram Tracr")
h2("7.4 Skenario Penggunaan Utama")
rich([("Skenario A. Pedagang di pasar tanpa sinyal. ", True, False),
      ("Bu Sari berjualan nasi di pasar dengan sinyal hilang timbul. Ia membuka Tracr dari layar "
       "utama telepon genggamnya; aplikasi terbuka seketika karena cangkangnya sudah tersimpan di "
       "perangkat. Ia mengetuk tiga produk pada layar kasir lalu menekan “Catat Jualan”; catatan "
       "langsung muncul dan spanduk memberitahu satu catatan menunggu sinkronisasi. Sore hari, "
       "saat terhubung Wi-Fi rumah, seluruh antrian terkirim berurutan.", False, False)])
rich([("Skenario B. Menagih utang pelanggan. ", True, False),
      ("Pak Budi membuka halaman Utang-Piutang yang tersusun per orang. Kartu “Wati” menunjukkan "
       "total Rp185.000 dengan keterangan “Lewat 4 hari”; ia membentangkannya, lalu mengetuk "
       "tombol pengingat, dan aplikasi membuka WhatsApp dengan pesan yang sudah tersusun.",
       False, False)])
rich([("Skenario C. Mengajukan modal. ", True, False),
      ("Setelah enam bulan mencatat, Bu Sari membuka Laba Rugi, memilih periode satu tahun, lalu "
       "menekan “Cetak / PDF”. Dokumen hasilnya ia lampirkan dalam pengajuan kredit usaha "
       "rakyat, bukti tertulis yang sebelumnya tidak pernah ia miliki.", False, False)])

# ============================================================== 8. ANTARMUKA
h1("8. DESAIN ANTARMUKA")
h2("8.1 Prinsip Rancangan")
numbered("Satu tindakan utama per layar: satu tombol berwarna paling menonjol, sisanya netral.")
numbered("Angka lebih besar daripada label, sebab pelaku usaha mencari nominal.")
numbered("Status kosong yang menawarkan langkah berikutnya, bukan layar hampa.")
numbered("Kata sehari-hari: “Pelanggan ngutang”, “Modal”, “Untung”, bukan “piutang”, “harga "
         "pokok penjualan”, “laba bersih”; navigasi utama di bawah agar terjangkau ibu jari.")

h2("8.2 Tangkapan Layar")
table([
    ["[G1: Halaman Masuk]", "[G2: Beranda]", "[G3: Formulir Catat Transaksi]"],
    ["[G4: Aktivitas & penyaring]", "[G5: Kasir Sederhana]", "[G6: Nota Penjualan]"],
    ["[G7: Utang-Piutang]", "[G8: Laba Rugi]", "[G9: Laporan & grafik]"],
], widths=[4.7, 4.6, 4.7], header=False)

h2("8.3 Keterangan Antarmuka per Halaman")
table([
    ["Halaman", "Rute", "Deskripsi antarmuka"],
    ["Beranda", "/", "Kartu kekayaan bersih, ringkasan arus kas bulan berjalan, grafik pengeluaran, dek kartu akun, dan aktivitas terakhir"],
    ["Akun & rinciannya", "/accounts", "Kartu per akun terbagi menjadi Harta dan Utang; halaman rincian memuat grafik saldo harian, buku besar, dan penyesuaian saldo"],
    ["Aktivitas & formulir transaksi", "/transactions", "Daftar per hari, panel penyaring lengkap, mode pilih-banyak; formulir modal memuat tiga tab jenis transaksi, kolom berkalkulator, kategori bertingkat, label, dan lampiran struk"],
    ["Kasir Sederhana", "/products", "Petak produk berikon; ketukan menambah ke keranjang; bilah keranjang menempel di bawah; lembar pembayaran menyerupai nota"],
    ["Utang-Piutang", "/debts", "Kartu per orang yang dapat dibentangkan, usia utang berbahasa manusia, tombol bayar dan pengingat WhatsApp per baris"],
    ["Laba Rugi & Laporan", "/profit, /reports", "Kartu Penjualan, Modal, Laba Kotor, Biaya, Laba Bersih; pemilih rentang tanggal, donat kategori, peta panas harian, cetak/PDF"],
    ["Perencanaan & Buku", "/budgets, /books", "Anggaran, tagihan, dan target tabungan dalam satu halaman bertab; daftar buku dengan lencana jenis usaha atau pribadi"],
], widths=[3.0, 2.8, 8.2])

h2("8.4 Aksesibilitas")
bullet("Bahasa Indonesia sebagai bahasa bawaan; pengaturan ukuran teks di dalam aplikasi, "
       "terpisah dari pengaturan sistem; mode terang dan gelap dengan rasio kontras yang dijaga; "
       "sasaran ketukan yang lapang dan kerangka pemuatan berukuran tetap sehingga tata letak "
       "tidak melompat saat data tiba.")
bullet("Yang paling menentukan: aplikasi tetap dapat dipakai penuh tanpa jaringan.")

# ================================================================= LAMPIRAN
page_break(h1("LAMPIRAN"))
table([
    ["Keterangan", "Tautan"],
    ["Website (aktif, dapat diakses publik)", "https://tracr-ai.vercel.app"],
    ["Repositori kode sumber (publik)", "https://github.com/astrorehan/FinancialTracker"],
], widths=[6.0, 8.0])
figure("02-alur-luring.png", "Gambar 4. Alur pencatatan luring dan sinkronisasi ulang "
       "(rujukan Bagian 7.1)", width=12.5)
h2("Daftar Pustaka")
for _ref in [
    "Badan Pusat Statistik. (2024). Statistik Potensi Desa Indonesia 2024 (Katalog 1105014, "
    "Nomor Publikasi 04300.24002). Jakarta: Badan Pusat Statistik. Diakses dari "
    "https://www.bps.go.id/id/publication/2024/12/10/2f5217e2d6a695a0830290a7/"
    "statistik-potensi-desa-indonesia-2024.html",

    "Fadhia, N., dan Ningsih, D. A. (2024). Penggunaan Pencatatan Akuntansi pada Usaha Mikro "
    "Kecil dan Menengah. Liabilities: Jurnal Pendidikan Akuntansi, 7(1), 30-37. ISSN 2620-5866. "
    "Diakses dari https://jurnal.umsu.ac.id/index.php/LIAB/article/view/15883",

    "Fatwitawati, R. (2018). Pengelolaan Keuangan bagi Usaha Mikro Kecil Menengah (UMKM) di "
    "Kelurahan Airputih, Kecamatan Tampan, Kota Pekanbaru. Sembadha: Seminar Hasil Pengabdian "
    "kepada Masyarakat. Jakarta: PKN STAN Press. Diakses dari "
    "https://jurnal.pknstan.ac.id/index.php/sembadha/article/view/376",

    "International Finance Corporation. (2025). MSME Finance Gap: An Updated Estimation and "
    "Evolution of the Micro, Small and Medium Enterprises Gap in Emerging and Developing "
    "Markets. Washington, DC: World Bank Group. Diakses dari "
    "https://www.smefinanceforum.org/data-sites/msme-finance-gap",

    "Kementerian Koperasi dan UKM. (2021). Data UMKM Semester I 2021, dikutip dalam Junaidi, M. "
    "(2024), UMKM Hebat, Perekonomian Nasional Meningkat. Direktorat Jenderal Perbendaharaan, "
    "Kementerian Keuangan Republik Indonesia. Diakses dari "
    "https://djpb.kemenkeu.go.id/kppn/curup/id/data-publikasi/artikel/"
    "2885-umkm-hebat,-perekonomian-nasional-meningkat.html",

    "Kementerian UMKM Republik Indonesia. (2025). Sistem Informasi Data Tunggal UMKM "
    "(SIDT-UMKM), posisi 31 Desember 2025. Diakses melalui https://satudata.kemenkopukm.go.id "
    "dan https://ukmindonesia.id/baca-deskripsi-posts/"
    "data-umkm-jumlah-dan-pertumbuhan-usaha-mikro-kecil-dan-menengah-di-indonesia",

    "Mayasari, M., Supardianto, Irawati, R., dan Lawita, N. F. (2025). Persepsi Pengguna "
    "terhadap Perangkat Lunak Akuntansi Usaha Kecil Berbasis Cloud. Jurnal Akuntansi dan "
    "Ekonomika, 15(1), 178-187. https://doi.org/10.37859/jae.v15i1.8505",

    "Otoritas Jasa Keuangan dan Badan Pusat Statistik. (2025). Siaran Pers Bersama: Indeks "
    "Literasi dan Inklusi Keuangan Masyarakat Meningkat, Hasil Survei Nasional Literasi dan "
    "Inklusi Keuangan (SNLIK) Tahun 2025. Nomor SP 69/OJK/GKPB/V/2025, 2 Mei 2025. Diakses dari "
    "https://ojk.go.id/id/berita-dan-kegiatan/siaran-pers/Pages/"
    "OJK-dan-BPS-Umumkan-Hasil-Survei-Nasional-Literasi-Dan-Inklusi-Keuangan-SNLIK-Tahun-2025.aspx",
]:
    _p = para(_ref, size=10, space_after=4, spacing=1.0)
    _p.paragraph_format.left_indent = Cm(1.0)
    _p.paragraph_format.first_line_indent = Cm(-1.0)

doc.save(OUT)
print("saved", OUT)
